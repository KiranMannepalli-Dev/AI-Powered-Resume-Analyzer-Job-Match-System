from docx import Document

def create_sample_resume():
    doc = Document()
    doc.add_heading('John Doe', 0)
    
    # Contact Info
    p = doc.add_paragraph()
    p.add_run('Email: john.doe@example.com | Phone: (555) 123-4567\n')
    p.add_run('LinkedIn: linkedin.com/in/johndoe | GitHub: github.com/johndoe')
    
    # Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph('Experienced Software Engineer with a strong background in Python, Flask, and React development. Passionate about building AI-powered applications and optimizing workflows.')
    
    # Skills
    doc.add_heading('Technical Skills', level=1)
    doc.add_paragraph('Programming: Python, JavaScript, SQL, C++\n'
                      'Web Technologies: Flask, React, Node.js, HTML, CSS\n'
                      'Cloud & DevOps: AWS, Docker, Git, Jenkins\n'
                      'Data Science: Pandas, NumPy, Scikit-learn')
    
    # Experience
    doc.add_heading('Work Experience', level=1)
    
    doc.add_heading('Senior Software Engineer | Tech Solutions Inc.', level=2)
    doc.add_paragraph('Jan 2021 - Present')
    doc.add_paragraph('- Developed and maintained scalable web applications using Flask and React.\n'
                      '- Implemented machine learning models to improve customer engagement by 25%.\n'
                      '- Led a team of 5 developers and improved deployment speed by 40% using Docker.')
    
    doc.add_heading('Junior Developer | StartUp Hub', level=2)
    doc.add_paragraph('June 2018 - Dec 2020')
    doc.add_paragraph('- Assisted in the development of RESTful APIs reaching 10k+ daily users.\n'
                      '- Optimized database queries resulting in 15% faster page load times.')
    
    # Education
    doc.add_heading('Education', level=1)
    doc.add_paragraph('Bachelor of Science in Computer Science | University of Technology | 2014 - 2018')
    
    doc.save('sample_resume.docx')
    print("Sample resume 'sample_resume.docx' created successfully.")

if __name__ == '__main__':
    create_sample_resume()
